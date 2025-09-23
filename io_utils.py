from docx import Document  # Word للتعامل مع ملفات
import pdfplumber  # PDF التي تستخرج النص من ملفات pdfplumber استيراد مكتبة


# (.txt) تابع لاستخراج النص من ملف نصي بسيط
def extract_txt(file_path: str):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


# (.docx) Word تابع لاستخراج النص من ملف
def extract_docx(file_path: str):
    doc = Document(file_path)
    # الموجودة في المستند (paragraphs) الحصول على قائمة بجميع الفقرات
    paras = [p.text for p in doc.paragraphs]
    # واحد وقائمة الفقرات بشكل منفصل string إرجاع النص المدمج في
    return "\n".join(paras), paras


# PDF تابع لاستخراج النص من ملف
def extract_pdf(file_path: str):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        # PDF المرور على كل صفحة في ملف
        for page in pdf.pages:
            # text استخراج النص من كل صفحة وإضافته إلى
            text += page.extract_text() or ""
    # تقسيم النص الكامل إلى فقرات وإرجاعه كقائمة.
    return text.split("\n")


# جديد وحفظه Word تابع لإنشاء ملف
def save_docx(paragraphs, output_path: str):
    doc = Document()
    # إضافة كل فقرة من القائمة إلى المستند الجديد
    for p in paragraphs:
        doc.add_paragraph(p)
    # حفظ المستند الجديد في المسار المحدد
    doc.save(output_path)
